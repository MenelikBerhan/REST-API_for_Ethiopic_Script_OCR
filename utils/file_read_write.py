"""Utilities for file reading & writing operations
"""
from config.setup import settings
from copy import deepcopy
from docx import Document
from fpdf import FPDF
from os import path, makedirs
from pdf2image import pdfinfo_from_bytes, convert_from_bytes
from PIL import Image
from typing import List, Tuple
from uuid import uuid4
import aiofiles
import io


async def background_write_file_image(file_buffer: bytes, file_name: str)\
        -> Tuple[Image.Image, dict]:
    """
    Retrieves image metadata from image buffer and saves image to local
    storage.

    Args:
        file_buffer (bytes): image file as bytes
        file_name (str): uploaded image's file name

    Returns:
        (tuple[PIL.Image, dict]): a tuple of image loaded as a PIL Image object
            and a dictionary containing image metadata.
    """
    # set storage path to `images` folder
    storage_path = path.join(settings.STORAGE_PATH, 'images')

    # check if storage path exists & is a directory. If not create one.
    if not (path.isdir(storage_path) or path.exists(storage_path)):
        print(f"Directory doesn't exist at given path: '{storage_path}'")
        makedirs(storage_path)
        print(f"Created storage directory: '{storage_path}")

    # set local file name by appending uuid value before the extension.
    base_name, ext = path.splitext(file_name)
    local_file_name = f'{base_name}_{uuid4()}{ext}'

    # set absolute local storage path by appending file name to storage path
    file_path = path.join(path.abspath(storage_path), local_file_name)

    # load image from bytes buffer using Pillow & BytesIO
    image = Image.open(io.BytesIO(file_buffer))

    # remove unnecessary info that causes errors for large png files
    info = deepcopy(image.info)
    info.pop('icc_profile', None)
    info.pop('exif', None)      # serialize it before use if needed

    # get image metadata and update image in db with it
    image_dict = {
        "local_path": file_path,
        "image_size": image.size,
        "image_format": image.format,
        "image_mode": image.mode,
        "info": info,
        }

    # save image to local storage asynchronously and return image & dict
    async with aiofiles.open(file_path, 'wb') as new_image_file:
        await new_image_file.write(file_buffer)

    return image, image_dict


async def background_write_file_pdf(file_buffer: bytes, file_name: str)\
        -> Tuple[List[Image.Image], dict]:
    """
    Retrieves pdf metadata from pdf buffer and saves pdf to local storage.

    Args:
        file_buffer (bytes): pdf file as bytes
        file_name (str): uploaded pdf's file name

    Returns:
        (tuple[list[PIL.Image], dict]): a tuple of a list of pdf's pages loaded
            as as PIL Image objects and a dictionary containing pdf metadata.
    """
    # set storage path to `pdfs` folder
    storage_path = path.join(settings.STORAGE_PATH, 'pdfs')

    # check if storage path exists & is a directory. If not create one.
    if not (path.isdir(storage_path) or path.exists(storage_path)):
        print(f"Directory doesn't exist at given path: '{storage_path}'")
        makedirs(storage_path)
        print(f"Created storage directory: '{storage_path}")

    # set local file name by appending uuid value before the extension (pdf).
    base_name, ext = path.splitext(file_name)
    local_file_name = f'{base_name}_{uuid4()}{ext}'

    # set absolute local storage path by appending file name to storage path
    file_path = path.join(path.abspath(storage_path), local_file_name)

    # get pdf info
    pdf_info_dict = pdfinfo_from_bytes(file_buffer)

    pdf_dict = {
        "local_path": file_path,
        'producer': pdf_info_dict['Producer'],
        'no_pages': pdf_info_dict['Pages'],
        'page_size': pdf_info_dict['Page size'],
        'file_size': pdf_info_dict['File size'],
        'pdf_version': pdf_info_dict['PDF version'],
    }

    # load pdf pages as PIL.Image objects
    pdf_images = convert_from_bytes(file_buffer, fmt='jpeg')

    # save pdf to local storage asynchronously and return pdf_images & dict
    async with aiofiles.open(file_path, 'wb') as new_pdf_file:
        await new_pdf_file.write(file_buffer)

    return pdf_images, pdf_dict


async def background_write_ocr_result_image(
        image_path: str, image_dict: dict, tess_output_dict: dict
        ):
    """
    Writes OCR result to files.

    Args:
        image_path (str): absolute path to image in local storage.
        image_dict (dict): dictionary dump of image in database
        tess_output_dict (dict): dictionary dump of OCR result in db
    """
    write_result_dict = {}

    # use image file name in local storage for ocr output file name.
    # Saves in same directory as image
    base_name, _ = path.splitext(image_path)
    text = tess_output_dict['ocr_result_text']

    # write OCR result to a plain text file if not done already
    if 'txt' in image_dict['ocr_output_formats'] and\
            'txt' not in image_dict['done_output_formats']:

        # set txt file path
        txt_file_path = base_name + '.txt'

        # write text to file & set path in write_result_dict
        async with aiofiles.open(txt_file_path, 'w') as txt_file:
            await txt_file.write(text)
        write_result_dict['txt'] = txt_file_path

    # write OCR result to a ms word (docx) file if not done already
    if 'docx' in image_dict['ocr_output_formats'] and\
            'docx' not in image_dict['done_output_formats']:
        # set docx file path
        docx_file_path = base_name + '.docx'

        # create word document and add text to new paragraph
        document = Document()
        par = document.add_paragraph().add_run(text)
        # set font (must use official font name)
        par.font.name = 'Abyssinica SIL'

        # save file and set path in write_result_dict
        document.save(docx_file_path)
        write_result_dict['docx'] = docx_file_path

    # write OCR result to a pdf file if not done already
    if 'pdf' in image_dict['ocr_output_formats'] and\
            'pdf' not in image_dict['done_output_formats']:
        # set pdf file path
        pdf_file_path = base_name + '.pdf'

        # get font path & name, line width & height
        font_path = './utils/fonts/AbyssinicaSIL-Regular.ttf'
        font_name = 'Abyssinica SIL'    # can use custom font names
        w, h = (0, 5)   # 0 width means use all available width

        # create pdf document & set params
        pdf = FPDF()
        # set params
        pdf.set_auto_page_break(True)
        pdf.add_font(font_name, fname=font_path)
        pdf.set_font(font_name)

        # add new page & write text
        pdf.add_page()
        pdf.multi_cell(w=w, h=h, text=text)

        # save pdf and set path in write_result_dict
        pdf.output(pdf_file_path)
        write_result_dict['pdf'] = pdf_file_path

    return write_result_dict


async def background_write_ocr_result_pdf(
        pdf_path: str, pdf_dict: dict, tess_output_dict: dict
        ):
    """
    Writes OCR result to files.

    Args:
        pdf_path (str): absolute path to pdf in local storage.
        pdf_dict (dict): dictionary dump of pdf in database
        tess_output_dict (dict): dictionary dump of OCR result in db
    """
    write_result_dict = {}

    # footer do demarcate each page in the pdf
    footer = '\n\t\t\t\t\t--- Page {} ---\n\n'

    # use pdf file name in local storage for ocr output file name.
    # Saves in same directory as pdf
    base_name, _ = path.splitext(pdf_path)
    text_by_page = tess_output_dict['ocr_result_text']

    # write OCR result to a plain text file if not done already
    if 'txt' in pdf_dict['ocr_output_formats'] and\
            'txt' not in pdf_dict['done_output_formats']:

        # set txt file path
        txt_file_path = base_name + '.txt'

        # concatenate texts from each page
        text = ''
        for page, page_text in text_by_page.items():
            text += page_text + footer.format(page)

        # write text to file & set path in write_result_dict
        async with aiofiles.open(txt_file_path, 'w') as txt_file:
            await txt_file.write(text)
        write_result_dict['txt'] = txt_file_path

    # write OCR result to a ms word (docx) file if not done already
    if 'docx' in pdf_dict['ocr_output_formats'] and\
            'docx' not in pdf_dict['done_output_formats']:
        # set docx file path
        docx_file_path = base_name + '.docx'

        # create word document and add text to new paragraph for each page
        document = Document()

        for page, page_text in text_by_page.items():
            # add footer to mark each page
            text = page_text + footer.format(page)

            par = document.add_paragraph().add_run(text)
            # set font (must use official font name)
            par.font.name = 'Abyssinica SIL'

            document.add_page_break()

        # save file and set path in write_result_dict
        document.save(docx_file_path)
        write_result_dict['docx'] = docx_file_path

    # write OCR result to a pdf file if not done already
    if 'pdf' in pdf_dict['ocr_output_formats'] and\
            'pdf' not in pdf_dict['done_output_formats']:
        # set pdf file path (add `-ocr` to distinguish from input pdf)
        pdf_file_path = base_name + '-ocr.pdf'

        # get font path & name, line width & height
        font_path = './utils/fonts/AbyssinicaSIL-Regular.ttf'
        font_name = 'Abyssinica SIL'    # can use custom font names
        w, h = (0, 5)   # 0 width means use all available width

        # create pdf document & set params
        pdf = FPDF()
        # set params
        pdf.set_auto_page_break(True)
        pdf.add_font(font_name, fname=font_path)
        pdf.set_font(font_name)

        # add new page & write text for each page of pdf ocr result
        for page, page_text in text_by_page.items():
            # add footer to mark each page
            text = page_text + footer.format(page)

            pdf.add_page()
            pdf.multi_cell(w=w, h=h, text=text)

        # save pdf and set path in write_result_dict
        pdf.output(pdf_file_path)
        write_result_dict['pdf'] = pdf_file_path

    return write_result_dict
